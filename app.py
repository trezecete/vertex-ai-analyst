from flask import Flask, request, jsonify, send_from_directory, send_file
from flask_cors import CORS
from bq_service import BigQueryService
from ai_service import AIService
from sheets_service import SheetsService
import os
import json
import io
from docx import Document

app = Flask(__name__, static_folder='static')
CORS(app)

@app.route('/')
def index():
    return send_from_directory(app.static_folder, 'index.html')

@app.route('/api/get_inventory', methods=['POST'])
def get_inventory():
    data = request.json
    project_id = data.get('project_id')
    dataset_id = data.get('dataset_id')
    sa_info = data.get('sa_info')

    if not project_id or not dataset_id:
        return jsonify({"error": "Project ID and Dataset ID are required."}), 400

    # New credentials logic
    creds = None
    if sa_info:
        creds = sa_info
    elif os.path.exists('key/key.json'):
        creds = 'key/key.json'

    try:
        bq = BigQueryService(project_id, service_account_info=creds)
        tables = bq.list_tables(dataset_id)
        
        inventory_summary = []
        for table in tables:
            # Pegamos apenas metadados básicos primeiro para ser rápido
            table_ref = bq.client.dataset(dataset_id).table(table.table_id)
            full_table = bq.client.get_table(table_ref)
            inventory_summary.append({
                "table_id": table.table_id,
                "column_count": len(full_table.schema)
            })

        return jsonify({"tables": inventory_summary})

    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/analyze', methods=['POST'])
def analyze():
    data = request.json
    project_id = data.get('project_id')
    dataset_id = data.get('dataset_id')
    location = data.get('location', 'us-central1')
    sa_info = data.get('sa_info')

    if not project_id or not dataset_id:
        return jsonify({"error": "Project ID and Dataset ID are required."}), 400

    # New credentials logic
    creds = None
    if sa_info:
        creds = sa_info
    elif os.path.exists('key/key.json'):
        creds = 'key/key.json'

    try:
        # 1. Fetch full metadata
        bq = BigQueryService(project_id, service_account_info=creds)
        inventory = bq.get_dataset_inventory(dataset_id)
        
        if not inventory:
            return jsonify({"error": "No tables found in the specified dataset."}), 404

        # 2. Analyze with AI
        ai = AIService(project_id, location=location, service_account_info=creds)
        analysis_report = ai.analyze_tables(inventory)

        return jsonify({
            "report": analysis_report,
            "table_count": len(inventory)
        })

    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/get_sheets_inventory', methods=['POST'])
def get_sheets_inventory():
    data = request.json
    urls = data.get('urls', [])
    sa_info = data.get('sa_info')

    if not urls:
        return jsonify({"error": "At least one URL is required."}), 400

    # Credentials logic
    creds = None
    if sa_info:
        creds = sa_info
    elif os.path.exists('key/key.json'):
        creds = 'key/key.json'

    try:
        sheets = SheetsService(service_account_info=creds)
        inventory = sheets.get_multiple_sheets_inventory(urls)
        
        inventory_summary = []
        for item in inventory:
            inventory_summary.append({
                "table_id": item["table_id"],
                "column_count": len(item["schema"])
            })

        return jsonify({"tables": inventory_summary})

    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/analyze_sheets', methods=['POST'])
def analyze_sheets():
    data = request.json
    project_id = data.get('project_id')
    urls = data.get('urls', [])
    location = data.get('location', 'us-central1')
    sa_info = data.get('sa_info')

    if not project_id or not urls:
        return jsonify({"error": "Project ID and URLs are required."}), 400

    # Credentials logic
    creds = None
    if sa_info:
        creds = sa_info
    elif os.path.exists('key/key.json'):
        creds = 'key/key.json'

    try:
        # 1. Fetch data from sheets
        sheets = SheetsService(service_account_info=creds)
        inventory = sheets.get_multiple_sheets_inventory(urls)
        
        if not inventory:
            return jsonify({"error": "No data found in the specified sheets."}), 404

        # 2. Analyze with AI
        ai = AIService(project_id, location=location, service_account_info=creds)
        analysis_report = ai.analyze_tables(inventory)

        return jsonify({
            "report": analysis_report,
            "table_count": len(inventory)
        })

    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/export_docx', methods=['POST'])
def export_docx():
    data = request.json
    report_text = data.get('report', '')

    if not report_text:
        return jsonify({"error": "No report content provided."}), 400

    try:
        document = Document()
        document.add_heading('Relatório de Análise de Dados', 0)

        # Simple markdown to docx conversion (paragraph by paragraph)
        lines = report_text.split('\n')
        for line in lines:
            if line.startswith('# '):
                document.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                document.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                document.add_heading(line[4:], level=3)
            elif line.strip():
                # Remove some markdown characters for cleaner text
                clean_line = line.replace('**', '').replace('__', '').replace('`', '')
                document.add_paragraph(clean_line)

        file_stream = io.BytesIO()
        document.save(file_stream)
        file_stream.seek(0)

        return send_file(
            file_stream,
            as_attachment=True,
            download_name='relatorio_analise.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/generate_dashboard', methods=['POST'])
def generate_dashboard():
    data = request.json
    project_id = data.get('project_id')
    dataset_id = data.get('dataset_id') # Para BQ
    urls = data.get('urls', []) # Para Sheets
    location = data.get('location', 'us-central1')
    sa_info = data.get('sa_info')

    creds = None
    if sa_info:
        creds = sa_info
    elif os.path.exists('key/key.json'):
        creds = 'key/key.json'

    try:
        inventory = []
        if dataset_id:
            bq = BigQueryService(project_id, service_account_info=creds)
            inventory = bq.get_dataset_inventory(dataset_id)
        elif urls:
            sheets = SheetsService(service_account_info=creds)
            inventory = sheets.get_multiple_sheets_inventory(urls)

        if not inventory:
            return jsonify({"error": "No data found to generate dashboard."}), 404

        ai = AIService(project_id, location=location, service_account_info=creds)
        config = ai.generate_dashboard_config(inventory)
        
        # Adicionar informações de origem para permitir refresh futuro
        config['source_info'] = {
            "project_id": project_id,
            "dataset_id": dataset_id,
            "urls": urls,
            "location": location
        }

        return jsonify(config)

    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/refresh_data', methods=['POST'])
def refresh_data():
    data = request.json
    source_info = data.get('source_info')
    sa_info = data.get('sa_info')

    if not source_info:
        return jsonify({"error": "No source info provided."}), 400

    project_id = source_info.get('project_id')
    dataset_id = source_info.get('dataset_id')
    urls = source_info.get('urls', [])
    location = source_info.get('location', 'us-central1')

    creds = None
    if sa_info:
        creds = sa_info
    elif os.path.exists('key/key.json'):
        creds = 'key/key.json'

    try:
        inventory = []
        if dataset_id:
            bq = BigQueryService(project_id, service_account_info=creds)
            inventory = bq.get_dataset_inventory(dataset_id)
        elif urls:
            sheets = SheetsService(service_account_info=creds)
            inventory = sheets.get_multiple_sheets_inventory(urls)

        return jsonify({"inventory": inventory})

    except Exception as e:
        return jsonify({"error": str(e)}), 500

if __name__ == '__main__':
    app.run(debug=True, port=5000)
